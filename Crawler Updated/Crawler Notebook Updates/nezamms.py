#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import json
import urllib.parse
from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError

# =========================
# Settings
# =========================
OUTPUT_FILENAME_TEMPLATE = "Nezams_IDs.{date}.json"
DOC_OUTPUT_DIR = "Nezams_Docs"
SITE_HOME = "https://nezams.com/"

# Remove these parts from the page before extracting text
UNWANTED_SELECTORS = [
    "div.fontsize.no-print",
    "span.share-icon",
    "span.total-readers",
    "div.subject-share",
    "span.numbe-s",
    "div#more-items",
    "ul#subject-nav-links",
]

# =========================
# Helpers
# =========================
def ensure_dir(path: str):
    Path(path).mkdir(parents=True, exist_ok=True)

def sanitize_filename(name: str) -> str:
    # Keep Arabic, letters, numbers, spaces, underscore, dash
    safe = re.sub(r"[^0-9A-Za-z\u0600-\u06FF _\-]+", "", name)
    trimmed = safe.strip() or "document"
    return trimmed[:180]

def make_paragraph_rtl(paragraph):
    """Force paragraph RTL in python-docx."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "true")
    pPr.append(bidi)

def save_docx(title: str, body: str, filename: str):
    doc = Document()
    section = doc.sections[0]
    # This flips page direction; some Word versions honor paragraph flags more reliably:
    section.right_to_left = True

    # Base style
    try:
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(14)
    except Exception:
        pass

    # Title (RTL + right aligned)
    p_title = doc.add_paragraph()
    p_title.paragraph_format.right_to_left = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_title.add_run(title)
    make_paragraph_rtl(p_title)

    # Body (RTL + right aligned)
    p_body = doc.add_paragraph()
    p_body.paragraph_format.right_to_left = True
    p_body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_body.add_run(body)
    make_paragraph_rtl(p_body)

    doc.save(filename)

# =========================
# Step 1: Fetch IDs JSON
# =========================
def fetch_and_save_ids_sync() -> str:
    """
    Opens the home page, waits for the first POST to admin-ajax.php,
    reads JSON, extracts items, and writes Nezams_IDs.{date}.json.
    Returns the JSON filename.
    """
    today = datetime.now().strftime("%m.%d.%Y")
    json_filename = OUTPUT_FILENAME_TEMPLATE.format(date=today)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True, args=["--no-sandbox", "--disable-dev-shm-usage"])
        page = browser.new_page()

        # Go to site
        page.goto(SITE_HOME, wait_until="domcontentloaded")

        # If the site needs a user action to trigger the AJAX call, do it here:
        # e.g., page.click("css=button#load-data")  # (not needed right now)

        # Wait for the first admin-ajax.php POST response
        def is_target(resp):
            return ("admin-ajax.php" in resp.url) and (resp.request.method.upper() == "POST")

        try:
            resp = page.wait_for_response(is_target, timeout=30_000)
        except PlaywrightTimeoutError:
            browser.close()
            raise RuntimeError("Timed out waiting for admin-ajax.php POST response.")

        # Parse response as JSON (fallback to text if needed)
        content_type = (resp.headers.get("content-type") or "").lower()
        if "application/json" in content_type:
            target_data = resp.json()
        else:
            # Some servers send text but it's still JSON
            try:
                target_data = json.loads(resp.text() or "{}")
            except Exception:
                target_data = {}

        # Try to extract nonce from POST body (WordPress usually sends it in the form body)
        req = resp.request
        post_data = req.post_data or ""
        form = urllib.parse.parse_qs(post_data)
        wpnonce = (form.get("_wpnonce") or [None])[0]

        # Build systems list
        systems = []
        for item in (target_data.get("data") or []):
            if item.get("id"):
                systems.append({
                    "id": item.get("id"),
                    "name": item.get("text"),
                    "url": item.get("link"),
                })

        # Save JSON
        with open(json_filename, "w", encoding="utf-8") as f:
            json.dump(systems, f, ensure_ascii=False, indent=2)

        print(f"Saved {len(systems)} entries to {json_filename}")
        if wpnonce:
            print(f"_wpnonce (from POST body): {wpnonce}")

        browser.close()

    return json_filename

# =========================
# Step 2: Visit URLs + Save DOCX
# =========================
def scrape_and_save_all_sync(json_filename: str):
    ensure_dir(DOC_OUTPUT_DIR)

    # Load the JSON file
    try:
        with open(json_filename, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Error: JSON file '{json_filename}' not found.")
        return
    except json.JSONDecodeError:
        print("Error: JSON file is not valid.")
        return

    items_to_scrape = [x for x in data if x.get("url")]
    print(f"Loaded {len(items_to_scrape)} URLs from: {json_filename}\n")

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(locale="ar-SA")
        page = context.new_page()

        for idx, item in enumerate(items_to_scrape, start=1):
            url = item["url"]
            item_id = item.get("id", "N/A")

            try:
                page.goto(url, timeout=30_000, wait_until="domcontentloaded")
                # Give dynamic parts a second to render
                page.wait_for_timeout(1500)

                html = page.content()
                soup = BeautifulSoup(html, "lxml")  # lxml is faster/cleaner if installed

                # Adjust these selectors if the site structure changes
                title_tag = soup.select_one("body > div.page > h1")
                content_div = soup.select_one("body > div.page > div.post-page > div")
                if not title_tag or not content_div:
                    print(f"⚠ Skipped (no content): ID {item_id} — {url}")
                    continue

                # Remove unwanted parts
                for css in UNWANTED_SELECTORS:
                    for tag in content_div.select(css):
                        tag.decompose()

                # Merge nested 'selectionShareable' spans (site-specific fix)
                for outer in content_div.select('span.selectionShareable[style="color: #993300;"]'):
                    inner_spans = outer.select('span.selectionShareable')
                    combined = ' '.join(s.get_text(strip=True) for s in inner_spans if s.get_text(strip=True))
                    if combined:
                        outer.string = combined
                    for s in inner_spans:
                        s.decompose()

                title = (title_tag.get_text(strip=True) or "بدون عنوان").strip()
                body_text = content_div.get_text(separator="\n", strip=True)

                safe_title = sanitize_filename(title.replace("/", "-").replace(":", "،"))
                out_path = os.path.join(DOC_OUTPUT_DIR, f"{safe_title}.docx")
                save_docx(title, body_text, out_path)

                print(f"✅ Saved {idx} — ID: {item_id} — {safe_title}")
            except Exception as e:
                print(f"❌ Failed {idx} — ID: {item_id} — {url} — {e}")

        browser.close()

# =========================
# Run both steps
# =========================
if __name__ == "__main__":
    # Step 1: fetch JSON
    json_file = fetch_and_save_ids_sync()

    # Step 2: read JSON and write DOCX files
    scrape_and_save_all_sync(json_file)

    print(f"\nAll done. DOCX files are in: {DOC_OUTPUT_DIR}\\")
